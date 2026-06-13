"""
METACOD GLP-1 module — Word (.docx) report renderer.

Produces a print-quality Microsoft Word document from a resolved
PharmacistAssessment (i18n already resolved via services.i18n.resolve_i18n).
Mirrors the section structure of reports.html_renderer but emits a native
.docx so a clinician can open, annotate, sign and archive the conclusion in
Word — the format requested for clinical sign-off and for the retroactive
Clalit batch (one .docx conclusion per patient).

Layers
  - "clinical" : default — what the physician sees / hands to the patient
  - "admin"    : adds rule_ids, mechanism traces, internal pattern terms and
                 literature sources (METACOD research / SaMD audit only)

Trilingual: section titles come from the small _LABELS table below; clinical
strings come pre-resolved in the assessment dict. Hebrew ("he") is laid out
right-to-left (paragraph bidi + right alignment).

Design palette is borrowed from the editorial METACOD cabinet UI
(python/ui/index.html) so a printed conclusion and the on-screen cabinet read
as one product.

Pure python-docx — no template file required. `pip install python-docx`.
"""

from __future__ import annotations

import io
from datetime import datetime
from typing import Any, Optional

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm


# ----- Editorial palette (hex, no leading '#') — mirrors python/ui/index.html -----

INK = "1A1815"
INK_SOFT = "4A4540"
INK_QUIET = "807A72"
PETROL = "1F4A4D"
PAPER_EDGE = "E5DFD3"
CREAM_DEEP = "F1ECE0"

_TRIAGE = {
    "green": {"fg": "2D6E4F", "bg": "E8F0E9"},
    "yellow": {"fg": "946A12", "bg": "FAF1D9"},
    "red": {"fg": "A8281B", "bg": "FBE7E2"},
}
_GRADE_FG = {"A": "2D6E4F", "B": "946A12", "C": "807A72"}
_SEVERITY_BAR = {"high": "A8281B", "moderate": "946A12", "low": "2D6E4F"}


# ----- Localised section titles & static labels -----

_LABELS: dict[str, dict[str, str]] = {
    "ru": {
        "title": "METACOD GLP-1 — Профилактика осложнений",
        "subtitle": "Клиническое заключение — система поддержки врачебных решений",
        "visit": "Визит",
        "patient": "Пациент",
        "age": "Возраст",
        "sex": "Пол",
        "generated": "Сформировано",
        "triage": "Триаж",
        "interactions": "Лекарственно-пациентные взаимодействия",
        "no_interactions": "Триггеров правил для этого визита нет.",
        "medication": "Препарат",
        "physician_actions": "Действия врача",
        "patient_education": "Памятка пациенту",
        "lab_plan": "План лабораторного контроля",
        "mechanism": "Механизм",
        "sources": "Источники",
        "forecast": "Клинический прогноз",
        "glp1_decision": "Решение по дозе GLP-1",
        "future_risks": "Будущие риски",
        "chronic_adjustments": "Коррекция хронической терапии",
        "dose_adjustments": "Индивидуализация дозы",
        "projected_labs": "Прогноз лабораторных показателей",
        "status": "Статус подписания",
        "rule_pack": "Пакет правил",
        "drug_db": "База препаратов",
        "disclaimer": (
            "Документ является поддержкой врачебных решений. Окончательное "
            "клиническое суждение лечащего врача имеет приоритет."
        ),
        "weeks_ahead": "нед.",
        "trend": "тренд",
    },
    "en": {
        "title": "METACOD GLP-1 — Complication Prevention",
        "subtitle": "Clinical conclusion — physician decision support",
        "visit": "Visit",
        "patient": "Patient",
        "age": "Age",
        "sex": "Sex",
        "generated": "Generated",
        "triage": "Triage",
        "interactions": "Drug–patient interactions",
        "no_interactions": "No rule triggers for this visit.",
        "medication": "Medication",
        "physician_actions": "Physician actions",
        "patient_education": "Patient education",
        "lab_plan": "Lab plan",
        "mechanism": "Mechanism",
        "sources": "Sources",
        "forecast": "Clinical forecast",
        "glp1_decision": "GLP-1 dosing decision",
        "future_risks": "Future risks",
        "chronic_adjustments": "Chronic medication adjustments",
        "dose_adjustments": "Dose individualization",
        "projected_labs": "Projected lab values",
        "status": "Sign-off status",
        "rule_pack": "Rule pack",
        "drug_db": "Drug DB",
        "disclaimer": (
            "This document is decision support; the treating physician's "
            "clinical judgment prevails."
        ),
        "weeks_ahead": "wk",
        "trend": "trend",
    },
    "he": {
        "title": "METACOD GLP-1 — מניעת סיבוכים",
        "subtitle": "מסקנה קלינית — תמיכה בהחלטות הרופא",
        "visit": "ביקור",
        "patient": "מטופל",
        "age": "גיל",
        "sex": "מין",
        "generated": "הופק",
        "triage": "טריאז'",
        "interactions": "אינטראקציות תרופה–מטופל",
        "no_interactions": "אין חוקים שהופעלו לביקור זה.",
        "medication": "תרופה",
        "physician_actions": "פעולות הרופא",
        "patient_education": "הנחיות למטופל",
        "lab_plan": "תוכנית מעבדה",
        "mechanism": "מנגנון",
        "sources": "מקורות",
        "forecast": "תחזית קלינית",
        "glp1_decision": "החלטת מינון GLP-1",
        "future_risks": "סיכונים עתידיים",
        "chronic_adjustments": "התאמת טיפול כרוני",
        "dose_adjustments": "התאמת מינון אישית",
        "projected_labs": "תחזית ערכי מעבדה",
        "status": "סטטוס חתימה",
        "rule_pack": "חבילת חוקים",
        "drug_db": "מאגר תרופות",
        "disclaimer": (
            "מסמך זה מהווה תמיכה בהחלטות; שיקול הדעת הקליני של הרופא המטפל גובר."
        ),
        "weeks_ahead": "שב'",
        "trend": "מגמה",
    },
}


def _label(locale: str, key: str) -> str:
    table = _LABELS.get(locale, _LABELS["ru"])
    return table.get(key, _LABELS["ru"].get(key, key))


# ----- Low-level docx helpers -----

def _shade(element, hex_fill: str) -> None:
    """Apply a background fill to a cell or paragraph (via its properties)."""
    pr = element.get_or_add_tcPr() if hasattr(element, "get_or_add_tcPr") else element
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    pr.append(shd)


def _shade_cell(cell, hex_fill: str) -> None:
    _shade(cell._tc.get_or_add_tcPr(), hex_fill)


def _shade_paragraph(paragraph, hex_fill: str) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    pPr.append(shd)


def _left_border(paragraph, hex_color: str, size_eighths: int = 24) -> None:
    """Coloured vertical bar on the leading edge — the 'rule card' accent."""
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    edge = OxmlElement("w:left")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(size_eighths))
    edge.set(qn("w:space"), "8")
    edge.set(qn("w:color"), hex_color)
    pbdr.append(edge)
    pPr.append(pbdr)


def _set_rtl(paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _run(paragraph, text: str, *, bold: bool = False, italic: bool = False,
         size: Optional[int] = None, color: Optional[str] = None):
    r = paragraph.add_run(str(text) if text is not None else "")
    r.bold = bold
    r.italic = italic
    if size is not None:
        r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = RGBColor.from_string(color)
    return r


class _Doc:
    """Thin builder wrapping a python-docx Document with locale/RTL awareness."""

    def __init__(self, locale: str, direction: str) -> None:
        self.locale = locale
        self.rtl = direction == "rtl"
        self.doc = Document()
        self._setup_page()

    def _setup_page(self) -> None:
        section = self.doc.sections[0]
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        style = self.doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.font.color.rgb = RGBColor.from_string(INK)

    def para(self, *, space_before: int = 0, space_after: int = 4):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        if self.rtl:
            _set_rtl(p)
        return p

    def heading(self, text: str) -> None:
        p = self.para(space_before=12, space_after=4)
        _run(p, text, bold=True, size=14, color=PETROL)
        # subtle underline rule
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), PAPER_EDGE)
        pbdr.append(bottom)
        pPr.append(pbdr)

    def subheading(self, text: str) -> None:
        p = self.para(space_before=8, space_after=2)
        _run(p, text, bold=True, size=11, color=INK_SOFT)

    def bullet(self, text: str, grade: Optional[str] = None) -> None:
        p = self.doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        if self.rtl:
            _set_rtl(p)
        _run(p, text)
        if grade:
            _grade_badge(p, grade)


def _grade_badge(paragraph, grade: str) -> None:
    g = str(grade or "").strip() or "C"
    color = _GRADE_FG.get(g, INK_QUIET)
    _run(paragraph, "  ")
    _run(paragraph, f"[{g}]", bold=True, size=8, color=color)


# ----- Section renderers -----

def _add_header(d: _Doc, visit_id: str, patient: dict, created_at: str) -> None:
    L = lambda k: _label(d.locale, k)

    title_p = d.para(space_after=0)
    _run(title_p, L("title"), bold=True, size=20, color=PETROL)
    sub_p = d.para(space_after=8)
    _run(sub_p, L("subtitle"), italic=True, size=10, color=INK_QUIET)

    # Meta table: 2 columns x rows
    rows = [
        (L("visit"), str(visit_id)),
        (L("patient"), str(patient.get("patient_id", ""))),
        (f'{L("age")} / {L("sex")}',
         f'{patient.get("age_years", "")} / {patient.get("sex", "")}'),
        (L("generated"), created_at),
    ]
    table = d.doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT if d.rtl else WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    for i, (k, v) in enumerate(rows):
        kc, vc = table.rows[i].cells
        _shade_cell(kc, CREAM_DEEP)
        kp = kc.paragraphs[0]
        vp = vc.paragraphs[0]
        if d.rtl:
            _set_rtl(kp)
            _set_rtl(vp)
        _run(kp, k, bold=True, size=9, color=INK_SOFT)
        _run(vp, v, size=9)


def _add_triage(d: _Doc, assessment: dict) -> None:
    L = lambda k: _label(d.locale, k)
    color = assessment.get("triage_color", "green")
    scheme = _TRIAGE.get(color, _TRIAGE["green"])
    text = (assessment.get("triage_summary_i18n_text")
            or assessment.get("triage_summary_i18n_key", ""))

    d.heading(L("triage"))
    p = d.para(space_before=2, space_after=6)
    _shade_paragraph(p, scheme["bg"])
    p.paragraph_format.left_indent = Pt(6)
    p.paragraph_format.right_indent = Pt(6)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    _run(p, text, bold=True, size=11, color=scheme["fg"])


def _add_interactions(d: _Doc, assessment: dict, layer: str) -> None:
    L = lambda k: _label(d.locale, k)
    d.heading(L("interactions"))

    interactions = assessment.get("interactions", [])
    if not interactions:
        p = d.para()
        _run(p, L("no_interactions"), italic=True, color=INK_QUIET)
        return

    for finding in interactions:
        severity = finding.get("severity", "moderate")
        grade = finding.get("evidence_grade", "C")
        rule_name = finding.get("rule_name", finding.get("rule_id", ""))
        med = finding.get("medication_name", "")

        head = d.para(space_before=8, space_after=2)
        _left_border(head, _SEVERITY_BAR.get(severity, INK_QUIET))
        head.paragraph_format.left_indent = Pt(8)
        _run(head, rule_name, bold=True, size=11, color=INK)
        _grade_badge(head, grade)

        medp = d.para(space_after=2)
        medp.paragraph_format.left_indent = Pt(8)
        _run(medp, f'{L("medication")}: ', bold=True, size=9, color=INK_SOFT)
        _run(medp, med, size=9)

        if layer == "admin":
            mech = finding.get("admin_trace", {}).get("mechanism", "")
            if mech:
                mp = d.para(space_after=2)
                mp.paragraph_format.left_indent = Pt(8)
                _run(mp, f'{L("mechanism")}: ', bold=True, size=8, color=INK_QUIET)
                _run(mp, mech, italic=True, size=8, color=INK_QUIET)

        _add_action_list(d, L("physician_actions"),
                         finding.get("physician_actions", []), with_grade=True,
                         default_grade=grade)
        _add_action_list(d, L("patient_education"),
                         finding.get("patient_actions", []))
        _add_lab_list(d, L("lab_plan"), finding.get("lab_plan", []))

        if layer == "admin":
            trace = finding.get("admin_trace", {})
            if trace:
                ap = d.para(space_after=2)
                ap.paragraph_format.left_indent = Pt(8)
                _shade_paragraph(ap, "F4F0FF")
                _run(ap, "ADMIN  ", bold=True, size=8, color="4A2E8A")
                _run(ap, f'rule_id={trace.get("rule_id", "")}  '
                         f'pattern={trace.get("pattern", "")}', size=8, color="4A2E8A")
                internal = trace.get("i18n_text_internal", "")
                if internal:
                    _run(ap, f'  {internal}', italic=True, size=8, color="4A2E8A")
            sources = finding.get("sources", [])
            if sources:
                sp = d.para(space_after=2)
                sp.paragraph_format.left_indent = Pt(8)
                _run(sp, f'{L("sources")}: ', bold=True, size=8, color=INK_QUIET)
                _run(sp, "; ".join(str(s) for s in sources), size=8, color=INK_QUIET)


def _add_action_list(d: _Doc, title: str, actions: list, *,
                     with_grade: bool = False, default_grade: str = "C") -> None:
    if not actions:
        return
    d.subheading(title)
    for a in actions:
        txt = a.get("i18n_text") or a.get("i18n_key", "")
        grade = a.get("evidence_grade", default_grade) if with_grade else None
        d.bullet(txt, grade=grade)


def _add_lab_list(d: _Doc, title: str, labs: list) -> None:
    if not labs:
        return
    d.subheading(title)
    for lab in labs:
        txt = lab.get("i18n_text") or lab.get("i18n_key", lab.get("test", ""))
        d.bullet(txt)


def _add_forecast(d: _Doc, assessment: dict, layer: str) -> None:
    L = lambda k: _label(d.locale, k)
    forecast = assessment.get("forecast", {})
    d.heading(L("forecast"))

    glp1 = forecast.get("glp1_dosing", {})
    if glp1:
        d.subheading(L("glp1_decision"))
        decision = glp1.get("decision", "")
        text = glp1.get("i18n_text") or glp1.get("i18n_key", "")
        p = d.para()
        _run(p, str(decision).upper(), bold=True, color=PETROL)
        _run(p, f' — {text}')
        _grade_badge(p, glp1.get("evidence_grade", "C"))
        for r in (glp1.get("rationale_i18n_texts") or glp1.get("rationale_i18n_keys", [])):
            d.bullet(r)

    future = forecast.get("future_risks", [])
    if future:
        d.subheading(L("future_risks"))
        for r in future:
            txt = r.get("i18n_text") or r.get("i18n_key", "")
            tf = r.get("timeframe_weeks", "")
            d.bullet(f'[+{tf} {L("weeks_ahead")}] {txt}', grade=r.get("evidence_grade"))

    adjustments = forecast.get("chronic_med_adjustments", [])
    if adjustments:
        d.subheading(L("chronic_adjustments"))
        for a in adjustments:
            med = a.get("medication_name", "")
            adj = a.get("adjustment_i18n_text") or a.get("adjustment_i18n_key", "")
            trig = a.get("trigger_i18n_text") or a.get("trigger_i18n_key", "")
            tf = a.get("timeframe_weeks", "")
            d.bullet(f'{med} — {adj} ({trig}, +{tf} {L("weeks_ahead")})',
                     grade=a.get("evidence_grade"))


def _add_v3_sections(d: _Doc, assessment: dict) -> None:
    L = lambda k: _label(d.locale, k)

    dose = assessment.get("dose_adjustments", [])
    if dose:
        d.heading(L("dose_adjustments"))
        for a in dose:
            name = a.get("medication_name", a.get("drug_profile_id", ""))
            atype = a.get("adjustment_type", "")
            mag = a.get("adjustment_magnitude_percent")
            rationale = a.get("rationale_i18n_text") or a.get("rationale_i18n_key", "")
            mag_s = f' {mag:+.0f}%' if isinstance(mag, (int, float)) else ""
            d.bullet(f'{name}: {atype}{mag_s} — {rationale}', grade=a.get("evidence_grade"))

    labs = assessment.get("projected_labs", [])
    if labs:
        d.heading(L("projected_labs"))
        for lab in labs:
            name = lab.get("test_name", "")
            val = lab.get("projected_value")
            lo = lab.get("ci_low")
            hi = lab.get("ci_high")
            date = lab.get("projected_date", "")
            trend = lab.get("trend_direction", "")
            ci = (f' (95% CI {lo:g}–{hi:g})'
                  if isinstance(lo, (int, float)) and isinstance(hi, (int, float)) else "")
            val_s = f'{val:g}' if isinstance(val, (int, float)) else str(val)
            d.bullet(f'{name}: {val_s}{ci} @ {date} — {L("trend")}: {trend}')


def _add_footer(d: _Doc, assessment: dict, sign_off_status: str) -> None:
    L = lambda k: _label(d.locale, k)
    p = d.para(space_before=14, space_after=2)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "6")
    top.set(qn("w:space"), "4")
    top.set(qn("w:color"), PAPER_EDGE)
    pbdr.append(top)
    pPr.append(pbdr)
    _run(p, f'{L("status")}: ', bold=True, size=9, color=INK_SOFT)
    _run(p, sign_off_status, size=9, color=PETROL)
    _run(p, f'   |   {L("rule_pack")}: {assessment.get("rule_pack_version", "")}',
         size=9, color=INK_QUIET)
    _run(p, f'   |   {L("drug_db")}: {assessment.get("drug_db_version", "")}',
         size=9, color=INK_QUIET)

    dp = d.para(space_after=0)
    _run(dp, L("disclaimer"), italic=True, size=8, color=INK_QUIET)


# ----- Public entry points -----

def build_report_docx(
    visit_id: str,
    patient: dict,
    assessment: dict,
    locale: str = "ru",
    direction: str = "ltr",
    layer: str = "clinical",
    sign_off_status: str = "pending_sign_off",
    created_at: Optional[str] = None,
) -> bytes:
    """Render a single conclusion to .docx and return the file bytes.

    Args mirror reports.html_renderer.render_report_html. `assessment` must be
    the PharmacistAssessment dumped with mode="json" and passed through
    services.i18n.resolve_i18n for the chosen locale.
    """
    created_at = created_at or datetime.utcnow().isoformat(timespec="seconds")

    d = _Doc(locale=locale, direction=direction)
    _add_header(d, visit_id, patient, created_at)
    _add_triage(d, assessment)
    _add_interactions(d, assessment, layer)
    _add_forecast(d, assessment, layer)
    _add_v3_sections(d, assessment)
    _add_footer(d, assessment, sign_off_status)

    buf = io.BytesIO()
    d.doc.save(buf)
    return buf.getvalue()


def save_report_docx(path: str, **kwargs) -> str:
    """Convenience wrapper: build the docx and write it to `path`."""
    data = build_report_docx(**kwargs)
    with open(path, "wb") as f:
        f.write(data)
    return path
